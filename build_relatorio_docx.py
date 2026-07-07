from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
SOURCE_MD = BASE_DIR / "docs" / "relatorio_consolidado_ajustes_e_especificacoes.md"
OUTPUT_DOCX = BASE_DIR / "docs" / "relatorio_consolidado_ajustes_e_especificacoes.docx"

ALECE_GREEN = "46794A"
ALECE_DARK_GREEN = "2F5634"
ALECE_GOLD = "CDA85B"
LIGHT_GREEN = "EAF3EA"
LIGHT_GRAY = "F2F4F7"
TEXT = "212529"


def fix_text(value: str) -> str:
    replacements = {
        "Relatorio": "Relatório",
        "especificacoes": "especificações",
        "analise": "análise",
        "Ceara": "Ceará",
        "publicos": "públicos",
        "exercicio": "exercício",
        "minimos": "mínimos",
        "compreensao cidadÃ£": "compreensão cidadã",
        "compreensao cidadă": "compreensão cidadã",
        "compreensao": "compreensão",
        "categorizacao": "categorização",
        "anulacoes": "anulações",
        "exportacao": "exportação",
        "apos": "após",
        "necessarios": "necessários",
        "Correcao": "Correção",
        "correcao": "correção",
        "funcao": "função",
        "numericos": "numéricos",
        "mes e periodo": "mês e período",
        "Padronizacao": "Padronização",
        "padronizacoes": "padronizações",
        "inconsistencias": "inconsistências",
        "abreviacoes": "abreviações",
        "pontuacao": "pontuação",
        "correcoes": "correções",
        "Conferencia": "Conferência",
        "diferenca": "diferença",
        "periodos": "períodos",
        "suplencias": "suplências",
        "substituicoes": "substituições",
        "diferencas": "diferenças",
        "composicao": "composição",
        "espacos": "espaços",
        "usuario": "usuário",
        "Preservacao": "Preservação",
        "cidadao": "cidadão",
        "nao": "não",
        "decisao": "decisão",
        "contabil": "contábil",
        "lancamentos": "lançamentos",
        "classificacao": "classificação",
        "descricao": "descrição",
        "Alimentacao": "Alimentação",
        "refeicao": "refeição",
        "juridica": "jurídica",
        "Divulgacao": "Divulgação",
        "servicos": "serviços",
        "graficos": "gráficos",
        "Locacao": "Locação",
        "veiculos": "veículos",
        "Solicitacao": "Solicitação",
        "interpretacao": "interpretação",
        "cabecalho": "cabeçalho",
        "Informacao": "Informação",
        "incluÃ­do": "incluído",
        "Ã©": "é",
        "Ã s": "às",
        "exercÃ­cio": "exercício",
        "Tambem": "Também",
        "incluidos": "incluídos",
        "nÂº": "nº",
        "marco": "março",
        "Resolucao": "Resolução",
        "rodape": "rodapé",
        "Transparencia": "Transparência",
        "Gestao": "Gestão",
        " ate 2021": " até 2021",
        "Administracao": "Administração",
        "Ultima": "Última",
        "atualizacao": "atualização",
        "Graficos": "Gráficos",
        "contem": "contém",
        "grafico": "gráfico",
        "evolucao": "evolução",
        "tecnicos": "técnicos",
        "amigaveis": "amigáveis",
        "exibicao": "exibição",
        "unica": "única",
        "opcao": "opção",
        "Endereco": "Endereço",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_inches: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ALECE_DARK_GREEN)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ALECE_DARK_GREEN, 16, 8),
        ("Heading 2", 13, ALECE_GREEN, 12, 6),
        ("Heading 3", 12, ALECE_DARK_GREEN, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Relatório consolidado")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ALECE_DARK_GREEN)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Ajustes da base e especificações do dashboard VDP ALECE 2025")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor.from_string(TEXT)

    p = doc.add_paragraph()
    p.add_run("Projeto: ").bold = True
    p.add_run("Painel da Verba de Desempenho Parlamentar")

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2.1, 4.2])
    data = [
        ("Base analisada", "Verba de Desempenho Parlamentar - Exercício 2025"),
        ("Registros", "4.219"),
        ("Despesa após anulações", "R$ 19.893.422,35"),
        ("Arquivo do dashboard", "dashboard-vdp-alece-atualizado.zip"),
    ]
    for row, (label, value) in zip(table.rows, data):
        set_cell_shading(row.cells[0], LIGHT_GREEN)
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    doc.add_page_break()


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    header = [fix_text(cell) for cell in split_table_row(rows[0])]
    body = [[fix_text(cell) for cell in split_table_row(row)] for row in rows[2:]]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [6.5 / len(header)] * len(header)
    if len(header) == 4:
        widths = [2.25, 2.25, 0.9, 1.1]
    elif len(header) == 2:
        widths = [2.25, 4.25]
    set_table_width(table, widths)

    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.runs[0].font.bold = True
            paragraph.paragraph_format.space_after = Pt(0)

    for values in body:
        row = table.add_row()
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            cell.text = text
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.05)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(fix_text(text))


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(fix_text(text))


def add_paragraph_with_links(doc: Document, line: str) -> None:
    p = doc.add_paragraph()
    line = fix_text(line)
    pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
    pos = 0
    for match in pattern.finditer(line):
        if match.start() > pos:
            p.add_run(line[pos : match.start()])
        add_hyperlink(p, match.group(1), match.group(2))
        pos = match.end()
    if pos < len(line):
        p.add_run(line[pos:])


def build_docx() -> None:
    text = fix_text(SOURCE_MD.read_text(encoding="utf-8"))
    doc = Document()
    apply_styles(doc)
    add_cover(doc)

    table_buffer: list[str] = []
    in_code_block = False

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            if line.strip():
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue

        if line.startswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

        if not line.strip():
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        elif re.match(r"^\d+\.\s", line):
            add_numbered(doc, re.sub(r"^\d+\.\s", "", line))
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(ALECE_DARK_GREEN)
        else:
            add_paragraph_with_links(doc, line)

    if table_buffer:
        add_markdown_table(doc, table_buffer)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Dashboard VDP ALECE 2025 | Relatório consolidado").font.size = Pt(9)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
